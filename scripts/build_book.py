"""Build the MAIOS book DOCX (신국판) from the markdown docs.

Structure: cover -> preface -> contents -> Part 1 (manual) -> Part 2 (paper)
-> appendix (deliverables). Figures rendered by extract_figures.py are
inserted at their paper positions.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIGS = DOCS / "figures"
OUT = DOCS / "MAIOS_책_v1.docx"

BODY_FONT = "바탕"
HEAD_FONT = "맑은 고딕"
CODE_FONT = "Consolas"
INK2 = RGBColor(0x52, 0x51, 0x4E)
BLUE = RGBColor(0x1C, 0x5C, 0xAB)

FIGURES = {
    "4.1": ("fig1.png", "그림 1. MAIOS 계층 아키텍처와 6단계 인지 루프. 학습 단계의 교훈은 점선 경로를 따라 다음 사이클에 환류된다."),
    "5.2": ("fig2.png", "그림 2. 스프린트별 누적 테스트 수(347→506). 분기 커버리지 95% 게이트를 유지한 채 이틀간 18개 지점에서 증가했다."),
    "6.2": ("fig3.png", "그림 3. 자기 기록 되먹임 팽창의 순환(적색)과 세 차단 지점(녹색)."),
    "6.3": ("fig4.png", "그림 4. 세 결함 수정의 정량 효과 (동일 실코퍼스, 수정 전 → 후)."),
    "7.1": ("fig5.png", "그림 5. 온톨로지 접합부 4곳. J3(점선)은 향후 과제."),
    "7.2": ("fig6.png", "그림 6. 의도 간극의 온톨로지 매개(7.2절 실증)."),
    "7.4": ("fig7.png", "그림 7. 의도-행동 정렬의 3치 판정 흐름."),
}


def set_font(run, name=BODY_FONT, size=10.5, bold=False, color=None, code=False):
    run.font.name = CODE_FONT if code else name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CODE_FONT if code else name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_runs(paragraph, text, size=10.5, base_font=BODY_FONT, color=None):
    """Emit text with **bold** and `code` inline markup."""
    for token in re.split(r"(\*\*.+?\*\*|`[^`]+`)", text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_font(run, base_font, size, bold=True, color=color)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, size=size - 1, code=True, color=color)
        else:
            run = paragraph.add_run(token)
            set_font(run, base_font, size, color=color)


def para(doc, text="", size=10.5, bold=False, align=None, font=BODY_FONT,
         color=None, space_after=6, line=1.5, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        if bold:
            run = p.add_run(re.sub(r"\*\*|`", "", text))
            set_font(run, font, size, bold=True, color=color)
        else:
            add_runs(p, text, size, font, color)
    return p


def heading(doc, text, level):
    sizes = {1: 16, 2: 13, 3: 11.5}
    p = para(doc, text, size=sizes[level], bold=True, font=HEAD_FONT,
             space_after=8, line=1.3)
    p.paragraph_format.space_before = Pt({1: 22, 2: 16, 3: 12}[level])
    p.paragraph_format.keep_with_next = True
    return p


def code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Cm(0.4)
        run = p.add_run(line if line else " ")
        set_font(run, size=8.5, code=True, color=INK2)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(8)


def md_table(doc, rows):
    header = [c.strip() for c in rows[0].strip("|").split("|")]
    body = [[c.strip() for c in r.strip("|").split("|")] for r in rows[2:]]
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell_paragraph = cell.paragraphs[0]
        add_runs(cell_paragraph, f"**{cell_text}**", size=9, base_font=HEAD_FONT)
    for i, row in enumerate(body, start=1):
        for j in range(len(header)):
            text = row[j] if j < len(row) else ""
            add_runs(table.rows[i].cells[j].paragraphs[0], text, size=9)
    para(doc, "", space_after=6)


def figure(doc, filename, caption):
    path = FIGS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.add_run().add_picture(str(path), width=Mm(112))
    cap = para(doc, caption, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER,
               font=HEAD_FONT, color=INK2, space_after=12, line=1.35)
    cap.paragraph_format.keep_together = True


def convert_markdown(doc, md_text, *, skip_first_heading=True, figures=None,
                     skip_code=False, demote=0):
    figures = figures or {}
    lines = md_text.splitlines()
    i = 0
    buffer: list[str] = []
    pending_figure = None
    first_heading = True

    def flush():
        nonlocal buffer
        if buffer:
            para(doc, " ".join(buffer))
            buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            flush()
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            if not skip_code:
                code_block(doc, block)
            i += 1
            continue
        if stripped.startswith("#"):
            flush()
            if pending_figure:
                figure(doc, *pending_figure)
                pending_figure = None
            level = min(3, len(stripped) - len(stripped.lstrip("#")) + demote)
            title = stripped.lstrip("#").strip()
            if first_heading and skip_first_heading:
                first_heading = False
                i += 1
                continue
            first_heading = False
            heading(doc, title, max(1, level))
            for key, fig in figures.items():
                if title.startswith(key):
                    pending_figure = fig
            i += 1
            continue
        if stripped.startswith("|"):
            flush()
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            if len(rows) >= 2:
                md_table(doc, rows)
            continue
        if re.match(r"^[-*] ", stripped) or re.match(r"^\d+\. ", stripped):
            flush()
            item = [re.sub(r"^([-*]|\d+\.) ", "", stripped)]
            i += 1
            while i < len(lines) and lines[i].startswith(("  ", "\t")) and lines[i].strip():
                item.append(lines[i].strip())
                i += 1
            marker = "• " if re.match(r"^[-*] ", stripped) else stripped.split(" ")[0] + " "
            p = para(doc, "", indent=0.4, space_after=3)
            add_runs(p, marker + " ".join(item))
            continue
        if stripped in {"---", "***"}:
            flush()
            i += 1
            continue
        if not stripped:
            flush()
            i += 1
            continue
        buffer.append(stripped)
        i += 1
    flush()
    if pending_figure:
        figure(doc, *pending_figure)


def page_number_footer(section):
    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer_paragraph._p.append(fld)


def part_page(doc, label, title, subtitle):
    doc.add_page_break()
    for _ in range(8):
        para(doc, "", space_after=0)
    para(doc, label, size=13, bold=True, font=HEAD_FONT, color=BLUE,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, title, size=22, bold=True, font=HEAD_FONT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, line=1.35)
    para(doc, subtitle, size=10.5, color=INK2, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(152)
    section.page_height = Mm(225)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(19)
    section.right_margin = Mm(19)
    page_number_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)

    # ---- cover
    for _ in range(6):
        para(doc, "", space_after=0)
    para(doc, "MAIOS", size=40, bold=True, font=HEAD_FONT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "기억하는 AI 운영체제", size=17, bold=True, font=HEAD_FONT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "설계 · 실증 · 사용법, 그리고 임무형지휘의 소프트웨어적 번역",
         size=11, color=INK2, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    for _ in range(10):
        para(doc, "", space_after=0)
    para(doc, "권영환 지음  ·  Claude(Anthropic) 공동 설계·구현", size=10.5,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "v1 준비판  ·  2026년 7월", size=9.5, color=INK2,
         align=WD_ALIGN_PARAGRAPH.CENTER)

    # ---- preface
    doc.add_page_break()
    heading(doc, "머리말", 1)
    for text in [
        "이 책은 하나의 질문에서 시작되었다. 대화가 끝나면 모든 것을 잊는 AI를, "
        "임무를 중심으로 기억하고 성찰하며 일하는 체계로 바꿀 수 있는가.",
        "MAIOS(MUSA AI Operating System)는 그 질문에 대한 작동하는 대답이다. "
        "관찰-이해-계획-행동-성찰-학습의 인지 루프, 실행 사이에도 지속되는 워크스페이스 기억, "
        "과거의 경험과 문서를 스스로 찾아 쓰는 회상, 그리고 고위험 행동 앞에서 인간의 승인을 "
        "기다리는 거버넌스. 여기에 군의 임무형지휘 온톨로지를 결합하여, \"의도를 공유하고 실행을 "
        "위임하되 한계를 명시한다\"는 지휘 이론이 소프트웨어 통제 구조로 번역될 수 있음을 보였다.",
        "책의 제1부는 사용자를 위한 안내서다. 설치부터 대화형 셸, 문서 흡수, 온톨로지 결합까지 "
        "모든 명령을 실제 출력과 함께 담았다. 제2부는 연구 기록이다. 설계 원칙과 아키텍처, 그리고 "
        "실사용 코퍼스가 드러낸 세 가지 규모 결함 — 특히 '자기 기록 되먹임 팽창' — 의 발견과 해소 "
        "과정을 정량적으로 보고한다. 부록에는 이틀간의 개발 이력 전체를 남겼다.",
        "이 책 자체가 본문에서 말하는 '인간 지휘, AI 실행'의 산물이다. 사람이 방향을 정하고 "
        "위험을 수용했으며, AI가 설계를 제안하고 구현하였다. 그 협업의 기록이 독자에게도 "
        "쓸모 있기를 바란다.",
        "2026년 7월, 권영환",
    ]:
        para(doc, text, space_after=10)

    # ---- contents (static)
    doc.add_page_break()
    heading(doc, "차례", 1)
    for entry, level in [
        ("머리말", 0),
        ("제1부  사용 안내", 0),
        ("MAIOS란 무엇인가 / 설치 / 5분 훑어보기", 1),
        ("명령어 상세 (pursue · project · research · ingest · align · shell)", 1),
        ("워크스페이스 / 거버넌스 / 실전 워크플로우", 1),
        ("파이썬 API / 문제 해결 / FAQ", 1),
        ("제2부  설계와 실증 (연구논문)", 0),
        ("서론 / 관련 연구 / 설계 원칙 / 아키텍처", 1),
        ("구현과 검증 방법론 / 실사용 코퍼스 사례연구", 1),
        ("온톨로지 결합과 임무형지휘의 소프트웨어적 번역", 1),
        ("한계와 향후 과제 / 결론 / 참고문헌", 1),
        ("부록  산출물과 개발 이력", 0),
    ]:
        para(doc, entry, size=10.5 if level else 11.5, bold=(level == 0),
             font=HEAD_FONT if level == 0 else BODY_FONT,
             indent=0.5 if level else None, space_after=4)

    # ---- part 1: manual
    part_page(doc, "제 1 부", "사용 안내",
              "설치부터 온톨로지 결합까지 — MAIOS 사용 설명서")
    manual = (DOCS / "MANUAL.ko.md").read_text(encoding="utf-8")
    convert_markdown(doc, manual)

    # ---- part 2: paper
    part_page(doc, "제 2 부", "설계와 실증",
              "연구논문: 기억·회상·거버넌스를 갖춘 임무 중심 인지형 AI 운영체제")
    paper = (DOCS / "PAPER.ko.md").read_text(encoding="utf-8")
    convert_markdown(doc, paper, figures=FIGURES, skip_code=True)

    # ---- appendix
    part_page(doc, "부록", "산출물과 개발 이력",
              "2026-07-09 ~ 07-10, 대화형 공동개발의 기록")
    inventory = (DOCS / "산출물목록.ko.md").read_text(encoding="utf-8")
    convert_markdown(doc, inventory)

    doc.save(OUT)
    print("saved:", OUT, OUT.stat().st_size, "bytes")


if __name__ == "__main__":
    build()
